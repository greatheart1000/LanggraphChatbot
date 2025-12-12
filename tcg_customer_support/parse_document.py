"""
解析 TCG 客服场景 flow.docx 文档
提取 14 大类场景及其子类
"""
import json
from typing import Dict, List

import docx


def _split_lines(text: str) -> List[str]:
    return [ln.strip() for ln in text.splitlines() if ln.strip()]


def _parse_cell_details(text: str) -> Dict:
    """
    从子类详情单元格中解析：
    - description: 原始文本
    - faq_samples: 以项目符号列出的示例问法
    - related_subcategories: "关联子类" 下的条目
    """
    lines = _split_lines(text)
    faq_samples: List[str] = []
    related: List[str] = []
    mode = "faq"

    for ln in lines:
        if "关联子类" in ln:
            mode = "related"
            continue
        is_bullet = ln.startswith(("•", "·", "-", "—", "▪", "*"))
        content = ln.lstrip("•·-—▪* ").strip()
        if mode == "related":
            if content:
                related.append(content)
        else:
            if is_bullet and content:
                faq_samples.append(content)

    return {
        "description": text.strip(),
        "faq_samples": faq_samples,
        "related_subcategories": related,
    }


def parse_docx(file_path: str) -> Dict:
    """
    解析 Word 文档，提取场景结构（支持表格结构的主类/子类/关联子类）
    
    返回结构：
    {
      category: {
        "description": "",
        "subcategories": [
          {
            "name": "",
            "description": "",
            "faq_samples": [...],
            "related_subcategories": [...]
          }
        ]
      }
    }
    """
    doc = docx.Document(file_path)

    scenarios: Dict[str, Dict] = {}
    # 优先解析表格，因为文档结构主要在表格中
    for table in doc.tables:
        current_category = None
        for row in table.rows:
            cells = row.cells
            if not cells:
                continue
            cat_text = cells[0].text.strip()
            sub_text = cells[1].text.strip() if len(cells) > 1 else ""
            detail_text = cells[2].text.strip() if len(cells) > 2 else ""

            if cat_text:
                current_category = cat_text
                scenarios.setdefault(current_category, {"description": "", "subcategories": []})

            if not current_category:
                continue

            if sub_text:
                # 去掉可能的编号前缀，如 "3: 子类名"
                sub_name = sub_text.split(":", 1)[-1].strip() if ":" in sub_text else sub_text
                details = _parse_cell_details(detail_text)
                scenarios[current_category]["subcategories"].append(
                    {
                        "name": sub_name,
                        "description": details["description"],
                        "faq_samples": details["faq_samples"],
                        "related_subcategories": details["related_subcategories"],
                    }
                )
            else:
                # 无子类名时，可能是分类描述
                if detail_text:
                    desc = scenarios[current_category].get("description", "")
                    desc = (desc + "\n" + detail_text).strip() if desc else detail_text
                    scenarios[current_category]["description"] = desc

    # 兜底：如果表格未覆盖全部内容，保留原有段落解析逻辑（仅类别）
    if not scenarios:
        current_category = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if "场景列出来" in text:
                continue
            if text not in scenarios:
                scenarios[text] = {"subcategories": [], "description": ""}
            current_category = text

    return scenarios

def extract_scenarios(file_path: str) -> Dict:
    """
    提取场景信息并保存为 JSON
    
    Args:
        file_path: Word 文档路径
        
    Returns:
        场景字典
    """
    try:
        scenarios = parse_docx(file_path)
        
        # 保存为 JSON
        output_path = file_path.replace('.docx', '_parsed.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(scenarios, f, ensure_ascii=False, indent=2)
        
        print(f"解析完成！共找到 {len(scenarios)} 大类场景")
        print(f"结果已保存到: {output_path}")
        
        return scenarios
    except Exception as e:
        print(f"解析错误: {e}")
        return {}

if __name__ == "__main__":
    import sys
    file_path = sys.argv[1] if len(sys.argv) > 1 else "../TCG 客服场景flow.docx"
    scenarios = extract_scenarios(file_path)
    
    # 打印摘要
    for category, info in scenarios.items():
        print(f"\n{category}:")
        print(f"  子类数量: {len(info['subcategories'])}")
        for subcat in info['subcategories']:
            print(f"    - {subcat['name']}")

